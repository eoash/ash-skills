"""
EO 계약서 docx 공용 서식 헬퍼.

`build_eo_*.py` 스크립트들이 공유하는 python-docx 유틸 — 표준 폰트(굴림/바탕),
여백, 조문 번호 체계, 표 음영·내부 여백 주입을 여기서 관리한다.

호환 정책: 이 모듈의 함수들은 build 스크립트가 생성하는 docx 바이트와
paragraph/table fingerprint가 동일하게 유지되어야 한다. 시그니처/디폴트
값 변경 시 `/tmp/deslop_baseline/fingerprint.py`로 회귀 검증 필수.
"""

from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_east_asia_font(run, font_name):
    """동아시아 폰트(한글) 명시 — rFonts의 w:eastAsia 속성"""
    run.font.name = font_name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)


def set_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                         first_line_indent=None, left_indent=None,
                         space_before=None, space_after=None,
                         line_spacing=1.23):
    pf = p.paragraph_format
    pf.alignment = align
    if first_line_indent is not None:
        pf.first_line_indent = first_line_indent
    if left_indent is not None:
        pf.left_indent = left_indent
    pf.space_before = space_before if space_before is not None else Pt(1)
    pf.space_after = space_after if space_after is not None else Pt(1)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = line_spacing


def add_title(doc, text, space_after=Pt(18)):
    """표지 상단 제목. `space_after`는 표지 레이아웃에 따라 조정 가능."""
    p = doc.add_paragraph()
    set_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.CENTER,
                         space_before=Pt(0), space_after=space_after)
    r = p.add_run(text)
    set_east_asia_font(r, "굴림")
    r.font.size = Pt(18)
    r.bold = True
    return p


def add_article_header(doc, text):
    """제 X 조【 제목 】"""
    p = doc.add_paragraph()
    set_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                         space_before=Pt(8), space_after=Pt(2))
    r = p.add_run(text)
    set_east_asia_font(r, "굴림")
    r.font.size = Pt(10)
    r.bold = True
    return p


def add_body(doc, text, indent_level=0):
    """
    indent_level 0: 일반 본문 (first_line_indent 0.8cm)
    indent_level 1: '1. 2.' 번호 항목 (left_indent 0.5cm)
    indent_level 2: '1) 2)' 하위 항목 (left_indent 1.0cm)
    """
    p = doc.add_paragraph()
    kwargs = dict(align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    if indent_level == 0:
        kwargs["first_line_indent"] = Cm(0.8)
    elif indent_level == 1:
        kwargs["left_indent"] = Cm(0.5)
    elif indent_level == 2:
        kwargs["left_indent"] = Cm(1.0)
    set_paragraph_format(p, **kwargs)
    r = p.add_run(text)
    set_east_asia_font(r, "바탕")
    r.font.size = Pt(10)
    return p


def add_subtitle(doc, text, bold=True, size=12, center=True):
    """섹션 소제목 (예: 별지 제1호 서식, 일반조건 등)"""
    p = doc.add_paragraph()
    align = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.JUSTIFY
    set_paragraph_format(p, align=align, space_before=Pt(14), space_after=Pt(10))
    r = p.add_run(text)
    set_east_asia_font(r, "굴림")
    r.font.size = Pt(size)
    r.bold = bold
    return p


def add_blank(doc, size=6):
    p = doc.add_paragraph()
    set_paragraph_format(p, space_before=Pt(0), space_after=Pt(0))
    r = p.add_run("")
    set_east_asia_font(r, "바탕")
    r.font.size = Pt(size)


def add_page_break(doc):
    p = doc.add_paragraph()
    r = p.add_run()
    r.add_break(WD_BREAK.PAGE)


def configure_default_style(doc, font_name="바탕", size=Pt(10)):
    """Normal 스타일의 폰트·크기를 EO 표준(바탕 10pt)으로 고정."""
    style = doc.styles["Normal"]
    style.font.name = font_name
    style.font.size = size
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)


def configure_eo_page(doc, top=Cm(3.0), bottom=Cm(2.0),
                      left=Cm(2.5), right=Cm(2.5)):
    """A4 + EO 표준 여백(상 3 / 하 2 / 좌우 2.5 cm)."""
    section = doc.sections[0]
    section.top_margin = top
    section.bottom_margin = bottom
    section.left_margin = left
    section.right_margin = right
    return section


def set_cell_shading(cell, hex_color):
    """셀 배경색 (w:shd). 연회색 라벨: 'F2F2F2'"""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_margins(cell, top=80, bottom=80, left=140, right=140):
    """셀 내부 여백 (twentieths of a point)"""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, value in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
        tcMar.append(node)
    tcPr.append(tcMar)
