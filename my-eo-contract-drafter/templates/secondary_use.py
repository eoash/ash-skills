"""
EO 표준 "클린본 2차 활용 라이선스 계약서" 템플릿.

기존 `scripts/contracts/build_eo_secondary_use_contract.py`(모두닥 v2, 2026-04-23)의
본문 구조를 그대로 파라미터화. 원 계약이 있는 영상의 "클린본"을 별도 계약으로
재라이선스하는 시나리오 — EO가 원본 저작권은 유지하면서 2차 편집 용도로
사용 허락을 부여.

회귀 기준: `examples/build_modoodoc.py`가 `/tmp/deslop_baseline/mododoc_baseline.json`
fingerprint와 100% 일치.
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL

from lib.eo_docx import (
    set_east_asia_font,
    set_paragraph_format,
    set_cell_shading,
    set_cell_margins,
    add_title,
    add_article_header,
    add_body,
    add_subtitle,
    add_page_break,
    configure_default_style,
    configure_eo_page,
)


def build_secondary_use_contract(
    *,
    partner: dict,
    eo: dict,
    # 원 계약 메타
    original_contract_date: str,      # 예: "2025년 8월 12일"
    original_project_name: str,        # 예: "모두닥 다큐멘터리"
    original_video_url: str,
    original_upload_date: str,         # 예: "2025년 10월 20일"
    # 라이선스 금액
    supply_price: int,
    korean_supply_price: str,
    korean_vat: str,
    korean_total: str,
    # 이용 조건
    usage_purpose: str,                # 예: '"파트너"의 채용광고 용도'
    usage_purpose_short: str,          # 예: '채용 및 인재 유치를 위한 광고·마케팅'
    usage_media: str,                  # 예: 'YouTube, Meta, TikTok, ...'
    duration: str,                     # 예: '본 계약 효력 발생일부터 기한의 정함이 없음 (무기한)'
    # 담당자
    eo_contact: dict,                  # {"name", "role", "email"}
    partner_contact: dict,             # {"name", "email"}
    # 내부 메모
    internal_memos: list,              # [(제목, 본문)]
    checklist: list,                   # [str]
    # 출력
    output_path: str,
    # 옵션: 라이선스 대상/제목 커스터마이즈
    contract_name: str = None,         # default: "{project} 영상 클린본 라이선스"
    license_target_desc: str = None,   # default: 표준 클린본 정의
) -> str:
    """
    클린본 2차 활용 라이선스 계약서 생성.

    원 계약 제7조 + 별지 "클린본" 정의에 근거하여 별도 체결되는 후속 계약.
    저작권은 EO 유지(양도 아님 — 이용허락 성격).
    """
    doc = Document()
    configure_eo_page(doc)
    configure_default_style(doc)

    # ═════════════ 표지 ═════════════
    title = contract_name or f"영상 클린본 라이선스 계약서"
    add_title(doc, title, space_after=Pt(6))

    # ── 상단 요약표 5행 2열 ──
    table = doc.add_table(rows=5, cols=2)
    table.style = "Table Grid"
    table.autofit = False
    for row in table.rows:
        row.cells[0].width = Cm(3.5)
        row.cells[1].width = Cm(12.5)

    def tcell(cell, text, is_label=False, size=10):
        cell.text = ""
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_margins(cell)
        if is_label:
            set_cell_shading(cell, "F2F2F2")
        p = cell.paragraphs[0]
        align = WD_ALIGN_PARAGRAPH.CENTER if is_label else WD_ALIGN_PARAGRAPH.LEFT
        set_paragraph_format(p, align=align,
                             space_before=Pt(2), space_after=Pt(2))
        r = p.add_run(text)
        set_east_asia_font(r, "굴림" if is_label else "바탕")
        r.font.size = Pt(size)
        r.bold = is_label

    def tcell_multi(cell, lines):
        cell.text = ""
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_margins(cell)
        for idx, line in enumerate(lines):
            p = cell.add_paragraph() if idx > 0 else cell.paragraphs[0]
            set_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.LEFT,
                                 space_before=Pt(1), space_after=Pt(1))
            r = p.add_run(line)
            set_east_asia_font(r, "바탕")
            r.font.size = Pt(10)

    # 행 0: 계약명
    tcell(table.cell(0, 0), "계약명", is_label=True)
    tcell(table.cell(0, 1), f"{original_project_name} 영상 클린본 라이선스")

    # 행 1: 라이선스 대상 (3줄)
    tcell(table.cell(1, 0), "라이선스 대상", is_label=True)
    tcell_multi(table.cell(1, 1), [
        f'원 저작물: {original_upload_date} EO 공식 YouTube 채널 게시 "{original_project_name}" 1편',
        f'영상 URL: {original_video_url}',
        '라이선스 대상: 원 저작물의 "클린본" (EO 워터마크·음악·자막이 제거된 버전) 1편',
    ])

    # 행 2: 계약금액
    tcell(table.cell(2, 0), "계약금액", is_label=True)
    tcell_multi(table.cell(2, 1), [
        f"공급가액:  일금 {korean_supply_price}원 (₩{supply_price:,})",
        f"세    액:  일금 {korean_vat}원 (₩{supply_price // 10:,})",
        f"합    계:  일금 {korean_total}원 (₩{supply_price + supply_price // 10:,})",
    ])

    # 행 3: 이용 기간
    tcell(table.cell(3, 0), "이용 기간", is_label=True)
    tcell(table.cell(3, 1), duration)

    # 행 4: 이용 목적
    tcell(table.cell(4, 0), "이용 목적", is_label=True)
    tcell(table.cell(4, 1), f'{usage_purpose} 숏폼 영상 제작 및 배포 (매체·지역·편수 제한 없음)')

    # ── 전문 ──
    p = doc.add_paragraph()
    set_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                         first_line_indent=Cm(0.8),
                         space_before=Pt(6), space_after=Pt(4))
    r = p.add_run(
        f'"{partner["name"]}"(이하 "파트너")와 "{eo["name"]}"(이하 "EO")는 '
        f'{original_contract_date} 체결된 "영상제작 및 홍보계약서"(이하 "원 계약") 제7조 제3항 '
        '및 [별지 제1호 서식] "클린본" 제2항에 근거하여, 원 계약의 "완성본" 클린본에 '
        '대한 이용허락에 관하여 다음과 같이 계약을 체결하고, 증거로서 각각 1통씩 보관한다.'
    )
    set_east_asia_font(r, "바탕")
    r.font.size = Pt(10)

    add_body(doc,
        "붙임서류: 1) 영상 클린본 라이선스 계약서 일반조건 1부,  "
        "2) [별지 제1호 서식] 이용허락 상세 명세 1부", indent_level=0)

    # 체결일
    p = doc.add_paragraph()
    set_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.CENTER,
                         space_before=Pt(8), space_after=Pt(6))
    r = p.add_run("2026년      월      일")
    set_east_asia_font(r, "바탕")
    r.font.size = Pt(11)

    # ── 서명란 6행 2열 ──
    sign_table = doc.add_table(rows=6, cols=2)
    sign_table.style = "Table Grid"
    sign_table.autofit = False
    for row in sign_table.rows:
        row.cells[0].width = Cm(8.0)
        row.cells[1].width = Cm(8.0)

    def sgn(cell, text, is_header=False, size=10):
        cell.text = ""
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_margins(cell, top=40, bottom=40, left=140, right=140)
        if is_header:
            set_cell_shading(cell, "F2F2F2")
        p = cell.paragraphs[0]
        align = WD_ALIGN_PARAGRAPH.CENTER if is_header else WD_ALIGN_PARAGRAPH.LEFT
        set_paragraph_format(p, align=align,
                             space_before=Pt(1), space_after=Pt(1))
        r = p.add_run(text)
        set_east_asia_font(r, "굴림" if is_header else "바탕")
        r.font.size = Pt(size)
        r.bold = is_header

    sgn(sign_table.cell(0, 0), "파트너", is_header=True, size=11)
    sgn(sign_table.cell(0, 1), "EO", is_header=True, size=11)
    sgn(sign_table.cell(1, 0), f"상호: {partner['name']}")
    sgn(sign_table.cell(1, 1), f"상호: {eo['name']}")
    sgn(sign_table.cell(2, 0), f"사업자등록번호: {partner['business_no']}")
    sgn(sign_table.cell(2, 1), f"사업자등록번호: {eo['business_no']}")
    sgn(sign_table.cell(3, 0), f"주소: {partner['address']}")
    sgn(sign_table.cell(3, 1), f"주소: {eo['address']}")
    sgn(sign_table.cell(4, 0), f"대표이사: {partner['ceo']}                          (인)")
    sgn(sign_table.cell(4, 1), f"대표이사: {eo['ceo']}                          (인)")
    sgn(sign_table.cell(5, 0), f"담당: {partner_contact['name']} / {partner_contact['email']}")
    sgn(sign_table.cell(5, 1),
        f"담당: {eo_contact['name']} {eo_contact['role']} / {eo_contact['email']}")

    add_page_break(doc)

    # ═════════════ 일반조건 10조 ═════════════
    add_subtitle(doc, "영상 클린본 라이선스 계약서 일반조건", size=13)

    p = doc.add_paragraph()
    set_paragraph_format(p, first_line_indent=Cm(0.8),
                         space_before=Pt(4), space_after=Pt(10))
    r = p.add_run(
        f'"{partner["name"]}"(이하 "파트너"라 한다)와 "{eo["name"]}"'
        f'(이하 "EO"라 한다)는 양 당사자 간 {original_contract_date} 체결된 "원 계약"의 '
        f'제7조 제3항에 따라, 원 계약의 최종 결과물인 "완성본"의 "클린본"에 대한 '
        f'이용허락에 관하여 다음과 같이 합의한다.'
    )
    set_east_asia_font(r, "바탕")
    r.font.size = Pt(10)

    _add_secondary_use_articles(
        doc, eo, supply_price, korean_supply_price, korean_vat, korean_total,
        original_project_name, original_video_url, original_upload_date,
        usage_purpose_short, usage_media, duration,
    )

    add_page_break(doc)

    # ═════════════ 별지 제1호 — 이용허락 상세 명세 ═════════════
    add_subtitle(doc, "[별지 제1호 서식] 이용허락 상세 명세", size=12)

    spec_items = [
        ('원 저작물',
         f'{original_upload_date} "EO"가 운영하는 공식 YouTube 채널에 게시된 "{original_project_name}" 1편\n'
         f'URL: {original_video_url}'),
        ('이용허락 대상',
         '위 "원 저작물"의 "클린본" — EO 워터마크, 음악, 자막이 제거된 버전의 영상 파일 1편'),
        ('이용 주체',
         f'"파트너" ({partner["name"]}) 및 "파트너"의 의뢰를 받아 채용광고를 집행·제작·송출하는 제3자'),
        ('이용 목적',
         f'{usage_purpose} ({usage_purpose_short})'),
        ('이용 방식',
         '"클린본"을 편집·가공하여 숏폼 등 "2차 저작물"을 제작하고, 이를 "파트너"가 운영·계약·집행하는 일체의 매체에 '
         '게시·송출·배포하는 행위'),
        ('매체', f'{usage_media}'),
        ('지역', '제한 없음'),
        ('편수', '제한 없음 (단, 원소스는 "클린본" 1편에 한정)'),
        ('이용 기간', duration),
        ('대가', f'공급가액 금 {korean_supply_price}원 (₩{supply_price:,}, VAT 별도) — 일시불'),
        ('납품물',
         '"클린본" 디지털 파일 1편 — 본 계약 체결일부터 7영업일 이내 "EO"가 지정하는 방식(클라우드 링크 등)으로 '
         '인도'),
        ('이용허락 범위 외 사항',
         '원 계약 [별지 제1호 서식]에 정의된 "추가편집본", "원본"은 본 계약의 이용허락 범위에 포함되지 아니하며, '
         '"파트너"가 원할 경우 "EO"와 별도 계약으로 구매한다.'),
    ]

    spec_table = doc.add_table(rows=len(spec_items), cols=2)
    spec_table.style = "Table Grid"
    spec_table.autofit = False
    for row in spec_table.rows:
        row.cells[0].width = Cm(4.0)
        row.cells[1].width = Cm(12.0)

    for idx, (label, desc) in enumerate(spec_items):
        tcell(spec_table.cell(idx, 0), label, is_label=True)
        value_lines = desc.split("\n")
        if len(value_lines) > 1:
            tcell_multi(spec_table.cell(idx, 1), value_lines)
        else:
            tcell(spec_table.cell(idx, 1), desc)

    # ═════════════ 내부 검토 메모 ═════════════
    add_page_break(doc)
    _add_internal_memo_page(doc, partner["name"], internal_memos, checklist)

    # 저장
    out = Path(output_path).expanduser()
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    print(f"saved: {out}  ({out.stat().st_size} bytes)")
    return str(out)


def _add_secondary_use_articles(doc, eo, supply_price, korean_supply_price,
                                korean_vat, korean_total,
                                original_project_name, original_video_url,
                                original_upload_date, usage_purpose_short,
                                usage_media, duration):
    """일반조건 10조 본문. 대부분 고정 문구."""
    # 제1조
    add_article_header(doc, "제 1 조 【 계약의 목적 】")
    add_body(doc,
        '본 계약의 목적은 "EO"가 "파트너"에게 원 계약의 "완성본"의 "클린본"을 '
        '"파트너"의 채용광고 용도로 편집·이용할 수 있도록 이용허락을 부여하는 데 '
        '있으며, 이에 따른 이용허락의 범위, 대가, 기간 및 양 당사자의 권리·의무를 '
        '규정함에 있다.')

    # 제2조
    add_article_header(doc, "제 2 조 【 용어의 정의 】")
    add_body(doc,
        '본 계약에서 사용하는 용어의 의미는 원 계약의 [별지 제1호 서식] "용어의 정의"를 '
        '준용하며, 추가로 다음과 같이 정의한다.', indent_level=1)
    add_body(doc,
        f'1) "원 저작물": 원 계약에 따라 "EO"가 제작하여 {original_upload_date} "EO"가 운영하는 '
        f'공식 YouTube 채널({original_video_url})에 게시한 "{original_project_name}" 1편',
        indent_level=2)
    add_body(doc,
        '2) "클린본": 원 계약 [별지 제1호 서식]에 정의된 바와 같이, "완성본"에서 EO 워터마크, '
        '음악, 자막이 제거된 버전의 영상 파일', indent_level=2)
    add_body(doc,
        '3) "2차 저작물": "파트너"가 본 계약에 따라 "클린본"을 편집·가공하여 제작하는 '
        '영상 (숏폼, 광고 소재 등 일체의 편집 결과물)', indent_level=2)

    # 제3조
    add_article_header(doc, "제 3 조 【 이용허락의 부여 및 범위 】")
    add_body(doc,
        '"EO"는 "파트너"에게 본 계약에 따라 "클린본"을 "파트너"의 채용광고 용도로 '
        '편집·이용할 수 있는 비독점적 이용권을 부여한다. 이용허락의 구체적 범위는 '
        '다음과 같다.', indent_level=1)
    add_body(doc, f'1) 이용 목적: "파트너"의 {usage_purpose_short}',
             indent_level=2)
    add_body(doc,
        '2) 이용 방식: "클린본"을 편집·가공하여 "2차 저작물"을 제작하고, 이를 '
        '"파트너"가 운영·계약·집행하는 일체의 매체에 게시·송출·배포',
        indent_level=2)
    add_body(doc, f'3) 매체: {usage_media}',
             indent_level=2)
    add_body(doc, '4) 지역: 제한 없음', indent_level=2)
    add_body(doc, '5) 편수: 제한 없음 (단, 원소스는 본 조 제2호의 "클린본" 1편에 한정)',
             indent_level=2)
    add_body(doc, f'6) 이용 기간: {duration}',
             indent_level=2)
    add_body(doc,
        '전항에도 불구하고 "파트너"는 원 계약 제3조 제7항 및 제8항에 따라 "EO"가 '
        '준수하는 공정거래위원회의 표시·광고 심사지침을 위반하는 방식으로 '
        '"클린본" 또는 "2차 저작물"을 이용하지 아니한다.', indent_level=1)

    # 제4조
    add_article_header(doc, "제 4 조 【 저작권의 귀속 】")
    add_body(doc,
        '"원 저작물" 및 "클린본"의 저작권과 이에 준하는 일체의 지식재산권은 원 계약 '
        '제6조 제1항에 따라 "EO"에 귀속되며, 본 계약은 이용허락에 관한 것으로서 '
        '저작권의 양도를 의미하지 아니한다.', indent_level=1)
    add_body(doc,
        '"파트너"가 본 계약에 따라 "클린본"을 편집·가공하여 제작한 "2차 저작물"의 '
        '저작권은 "파트너"에 귀속된다. 다만, "2차 저작물"에 포함된 "클린본" 원소스에 '
        '대한 "EO"의 저작권은 영향을 받지 아니한다.', indent_level=1)
    add_body(doc,
        '"파트너"는 "클린본" 원소스 자체를 본 계약의 이용 목적을 벗어나 제3자에게 '
        '양도·대여·재라이선스하거나, 독립된 영상 저작물로서 배포·판매할 수 없다.',
        indent_level=1)

    # 제5조
    add_article_header(doc, "제 5 조 【 라이선스료 및 지급 】")
    add_body(doc,
        f'"파트너"는 본 계약에 따른 이용허락의 대가로 "EO"에게 공급가액 금 {korean_supply_price}원'
        f'(₩{supply_price:,}, VAT 별도)을 지급한다. 부가가치세 금 {korean_vat}원(₩{supply_price // 10:,})을 '
        f'합한 총 지급액은 금 {korean_total}원(₩{supply_price + supply_price // 10:,})이다.',
        indent_level=1)
    add_body(doc,
        '"파트너"는 본 계약 체결일 및 "EO"의 "클린본" 인도 이후 14영업일 이내에 '
        '전항의 총 지급액을 일시불로 "EO"가 지정한 계좌로 지급한다.', indent_level=1)
    add_body(doc,
        '"EO"는 라이선스료 수령을 위하여 "파트너"에게 세금계산서를 발행한다.',
        indent_level=1)
    add_body(doc,
        f'입금 정보 — 은행: {eo["bank"]["bank_name"]} / '
        f'계좌번호: {eo["bank"]["account"]} / '
        f'예금주: {eo["bank"]["holder"]}', indent_level=1)

    # 제6조
    add_article_header(doc, "제 6 조 【 클린본의 인도 】")
    add_body(doc,
        '"EO"는 본 계약 체결일부터 7영업일 이내에 "EO"가 지정하는 방식(클라우드 '
        '링크 등)으로 "클린본" 디지털 파일을 "파트너"에게 인도한다.', indent_level=1)
    add_body(doc,
        '"파트너"는 "클린본" 파일을 선량한 관리자의 주의로 보관·관리하며, '
        '제3자에 의한 무단 복제·유출을 방지하기 위한 합리적 보안 조치를 취한다.',
        indent_level=1)

    # 제7조
    add_article_header(doc, "제 7 조 【 재라이선스 및 재양도 금지 】")
    add_body(doc,
        '"파트너"는 본 계약에 따라 부여받은 이용권 또는 "클린본" 원소스를 "EO"의 '
        '사전 서면 동의 없이 제3자에게 재라이선스하거나 양도·대여·담보제공 기타 '
        '처분할 수 없다. 다만, "2차 저작물"을 "파트너"의 채용광고 목적으로 '
        '광고대행사·매체사·프로덕션 등에 송출·집행을 의뢰하기 위하여 필요한 범위에서 '
        '전달하는 경우에는 그러하지 아니한다.')

    # 제8조
    add_article_header(doc, "제 8 조 【 원 계약 조항의 준용 】")
    add_body(doc,
        '본 계약에서 따로 정하지 아니한 사항에 관하여는 원 계약의 다음 각 호 조항을 '
        '그 성질에 반하지 아니하는 범위에서 준용한다.', indent_level=1)
    add_body(doc, '1) 원 계약 제8조 (비밀유지의무)', indent_level=2)
    add_body(doc, '2) 원 계약 제9조 (계약의 해지)', indent_level=2)
    add_body(doc, '3) 원 계약 제10조 (손해배상)', indent_level=2)
    add_body(doc, '4) 원 계약 제11조 (양도금지)', indent_level=2)
    add_body(doc, '5) 원 계약 제12조 (분쟁의 해결)', indent_level=2)
    add_body(doc, '6) 원 계약 제13조 (관계법령 등의 준용)', indent_level=2)
    add_body(doc,
        f'양 당사자가 본 계약과 관련하여 상대방에게 부담하는 손해배상 책임의 총액은 '
        f'본 계약상 공급가액(₩{supply_price:,})을 한도로 한다. 다만, 고의 또는 중대한 과실, '
        f'제7조(재라이선스 및 재양도 금지) 위반, 또는 준용되는 원 계약 제8조'
        f'(비밀유지의무) 위반의 경우에는 본 한도가 적용되지 아니한다.', indent_level=1)

    # 제9조
    add_article_header(doc, "제 9 조 【 계약의 효력 】")
    add_body(doc,
        '본 계약은 양 당사자가 서명·날인한 날부터 유효하며, 본 조 및 원 계약의 '
        '이행 완료 여부에 영향을 받지 아니한다.', indent_level=1)
    add_body(doc,
        '본 계약에 따라 "파트너"에게 부여된 이용권은 제5조에 정한 라이선스료 '
        '총액이 전액 지급되는 시점부터 행사할 수 있다.', indent_level=1)

    # 제10조
    add_article_header(doc, "제 10 조 【 기 타 】")
    add_body(doc,
        '본 계약의 변경 또는 추가는 양 당사자의 서면 합의에 의해서만 효력이 '
        '발생한다. 본 계약을 증명하기 위해 계약서 2부를 작성하고 각각 1부씩 '
        '보관한다.')


def _add_internal_memo_page(doc, partner_short_name, memos, checklist):
    """내부 검토 메모 — 외부 발송 전 삭제 필수."""
    p = doc.add_paragraph()
    set_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.CENTER,
                         space_before=Pt(6), space_after=Pt(10))
    r = p.add_run(f"⚠️ 내부 검토 메모 — {partner_short_name} 발송 전 이 페이지 반드시 삭제 ⚠️")
    set_east_asia_font(r, "굴림")
    r.font.size = Pt(12)
    r.bold = True
    r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    p = doc.add_paragraph()
    set_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.CENTER,
                         space_before=Pt(0), space_after=Pt(14))
    r = p.add_run(
        f"(본 페이지는 서현님 내부 검토용입니다. "
        f"{'모두사인' if '모두닥' in partner_short_name else partner_short_name}에 업로드되는 최종본에는 포함되지 않아야 합니다.)"
    )
    set_east_asia_font(r, "바탕")
    r.font.size = Pt(9)
    r.italic = True

    add_subtitle(doc, "서명 전 확인 필요 사항", size=12, center=False)

    for title, desc in memos:
        p = doc.add_paragraph()
        set_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                             space_before=Pt(8), space_after=Pt(2))
        r = p.add_run(title)
        set_east_asia_font(r, "굴림")
        r.font.size = Pt(10)
        r.bold = True

        p2 = doc.add_paragraph()
        set_paragraph_format(p2, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                             left_indent=Cm(0.5),
                             space_before=Pt(1), space_after=Pt(4))
        r2 = p2.add_run(desc)
        set_east_asia_font(r2, "바탕")
        r2.font.size = Pt(10)

    add_subtitle(doc, f"{partner_short_name} 발송 전 체크리스트", size=11, center=False)
    for item in checklist:
        p = doc.add_paragraph()
        set_paragraph_format(p, left_indent=Cm(0.3),
                             space_before=Pt(2), space_after=Pt(2))
        r = p.add_run(item)
        set_east_asia_font(r, "바탕")
        r.font.size = Pt(10)
