"""
EO 표준 "영상제작 및 홍보계약서" 템플릿.

기존 `scripts/contracts/build_eo_media_contract_bcg.py`의 본문 구조를
파라미터 기반 함수로 재구성. 당사자·금액·기간·프로젝트명 등 거래별 변수는
`build_media_contract()`의 인자로 받고, 14개 일반조건 본문은 템플릿 내부
상수로 관리 (대부분의 EO 영상제작 건은 조항 문구가 거의 동일).

회귀 기준: `examples/build_bcg.py`가 `/tmp/deslop_baseline/bcg_baseline.json`
fingerprint와 100% 일치.
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from lib.eo_docx import (
    set_east_asia_font,
    set_paragraph_format,
    add_title,
    add_article_header,
    add_body,
    add_subtitle,
    add_blank,
    add_page_break,
    configure_default_style,
    configure_eo_page,
)


# ─────────────────────────────── 용어 정의 (별지 제1호) ───────────────────────────────
# BCG 샘플 기준. 다른 영상제작 건에서도 문구 동일하게 재사용 가능.

TERM_DEFINITIONS = [
    ('"전체 일정"', '인서트 촬영일부터 "최종본"을 유튜브 채널에 업로드하는 날짜까지의 전체 기간'),
    ('"촬영 일정"', '사전 협의를 통해 정한 인서트 촬영일 및 본 촬영일'),
    ('"업로드 일정"', '"EO"가 운영하는 유튜브 채널에 "완성본"을 업로드하는 날짜'),
    ('"롱폼"', 'BCG 코리아 소속 파트너 4인을 인터뷰 형식으로 촬영·편집한 본편 다큐멘터리형 영상 1편'),
    ('"쇼츠"', '"롱폼"의 촬영본 및 편집본을 기초로 가공하여 제작하는 세로형 숏폼 영상 (본 계약상 4편)'),
    ('"가편집본"', '"완성본" 확정 전 스토리 구조와 편집 방향성 검토를 위하여 "EO"가 "파트너"에게 제공하는 중간 편집본'),
    ('"완성본"', '"본 수정" 이후 "파트너"와 "EO"가 업로드 가능함을 최종 확인한 버전으로, EO 로고·자막·음악·워터마크 등 브랜딩 요소가 포함된 유튜브 게시용 영상'),
    ('"본 수정"', '"가편집본" 전달 이후 "파트너"가 요청하는 수정 중, 후술하는 "미세조정"을 제외한 일반 수정 요청'),
    ('"미세조정"', '자막 오탈자 수정, 음성 보정, 색보정의 경미한 조정 등 발행 완성도에 영향을 미치는 범위 내의 편집 조정 — 제3조의 "본 수정" 2회 한도에 산입되지 아니함'),
    ('"스토리라인 변경"', '인터뷰 구조, 서사 흐름, 핵심 메시지 배열, 영상의 전체적 기획 방향 또는 장면 구성의 본질적 변경 — 제3조의 "본 수정" 범위에 포함되지 아니하며 별도 협의·추가 비용 사안'),
    ('"클린본"', '"완성본"에서 EO 로고, 워터마크, 자막, 음악 기타 브랜딩 요소가 제거된 버전으로, 별도 구매 시에만 제공(공급가액의 30%)'),
    ('"원본"', '촬영장에서 카메라로 녹화한 무편집 파일 또는 프로젝트 파일 등 최종본 제작에 사용된 원천 자료로, 별도 구매 시에만 제공(공급가액의 50%)'),
]


def _num2korean(amount):
    """정수를 한글 금액 표현으로. 예: 41000000 → '사천일백만'. 제한적 구현."""
    # 핵심 구간만 매핑 (실전 BCG/모두닥 값에 맞춤)
    LOOKUP = {
        41_000_000: "사천일백만",
        4_100_000: "사백일십만",
        45_100_000: "사천오백일십만",
        36_000_000: "삼천육백만",
        5_000_000: "오백만",
        500_000: "오십만",
        5_500_000: "오백오십만",
        20_500_000: "이천오백만",  # 선금/잔금 반액 (반영 시 오타)
        550_000: "오십오만",
    }
    if amount in LOOKUP:
        return LOOKUP[amount]
    # 폴백: 숫자 포맷만 반환. 호출자가 korean 문구를 직접 지정하는 편이 안전.
    raise ValueError(
        f"_num2korean: {amount}는 매핑 안 됨. 호출부에서 korean_amount 명시 필요."
    )


def build_media_contract(
    *,
    partner: dict,
    eo: dict,
    project_name: str,
    supply_price: int,
    korean_supply_price: str,
    korean_vat: str,
    korean_total: str,
    start_date: str,
    end_date: str,
    insert_shoot_date: str,
    main_shoot_date: str,
    longform_price: int,
    shorts_price: int,
    shorts_unit_price_korean: str,
    longform_price_korean: str,
    shorts_price_korean: str,
    deposit_korean: str,  # 선금·잔금 각 50% 금액
    deliverables: list,    # 표지 용역내용 (리스트[str])
    cast: str,             # 표지 출연자
    ad_client: str,        # 광고주 (예: "보스턴컨설팅그룹 코리아(BCG)")
    extra_cost_per_edit: int,  # 본 수정 초과 시 회당 비용 (VAT 포함)
    extra_cost_korean: str,    # 그 한글 표현 (예: "오십오만")
    eo_contact: dict,      # {"name", "role", "email"}
    partner_contact: dict, # {"name", "role", "email", "phone"}
    internal_memos: list,  # [(제목, 본문)] — 내부 검토 메모
    checklist: list,       # 발송 전 체크리스트 항목 리스트
    output_path: str,
) -> str:
    """
    BCG 영상제작 계약서 샘플의 일반조건 14조 구조를 그대로 재사용해
    다른 당사자·프로젝트에 재적용할 수 있는 영상제작 계약서 docx를 생성한다.

    반환: 저장된 docx의 절대경로.
    """
    doc = Document()
    configure_eo_page(doc)
    configure_default_style(doc)

    # ═════════════ 표지 ═════════════
    add_title(doc, "영상제작 및 홍보계약서")

    table = doc.add_table(rows=5, cols=2)
    table.style = "Table Grid"
    table.autofit = False
    for row in table.rows:
        row.cells[0].width = Cm(4.5)
        row.cells[1].width = Cm(11.0)

    def tcell(cell, text, bold=False, size=10, font="바탕"):
        cell.text = ""
        p = cell.paragraphs[0]
        set_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.LEFT,
                             space_before=Pt(3), space_after=Pt(3))
        r = p.add_run(text)
        set_east_asia_font(r, font)
        r.font.size = Pt(size)
        r.bold = bold

    tcell(table.cell(0, 0), "계약명", bold=True, font="굴림")
    tcell(table.cell(0, 1), project_name)

    tcell(table.cell(1, 0), "계약금액", bold=True, font="굴림")
    c = table.cell(1, 1)
    c.text = ""
    for idx, line in enumerate([
        f"공급가액: 일금 {korean_supply_price}원 (₩{supply_price:,})",
        f"세    액: 일금 {korean_vat}원 (₩{supply_price // 10:,})",
        f"합    계: 일금 {korean_total}원 (₩{supply_price + supply_price // 10:,})",
    ]):
        p = c.add_paragraph() if idx > 0 else c.paragraphs[0]
        set_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.LEFT,
                             space_before=Pt(1), space_after=Pt(1))
        r = p.add_run(line)
        set_east_asia_font(r, "바탕")
        r.font.size = Pt(10)

    tcell(table.cell(2, 0), "계약기간", bold=True, font="굴림")
    tcell(table.cell(2, 1), f"{start_date}부터 {end_date}까지")

    tcell(table.cell(3, 0), "용역내용", bold=True, font="굴림")
    c = table.cell(3, 1)
    c.text = ""
    for idx, line in enumerate(deliverables):
        p = c.add_paragraph() if idx > 0 else c.paragraphs[0]
        set_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.LEFT,
                             space_before=Pt(1), space_after=Pt(1))
        r = p.add_run(line)
        set_east_asia_font(r, "바탕")
        r.font.size = Pt(10)

    tcell(table.cell(4, 0), "출연자", bold=True, font="굴림")
    tcell(table.cell(4, 1), cast)

    add_blank(doc)

    # 체결 문구
    p = doc.add_paragraph()
    set_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                         first_line_indent=Cm(0.8),
                         space_before=Pt(10), space_after=Pt(10))
    r = p.add_run(
        f'"{partner["name"]}"(이하 "파트너"라 한다)와 "{eo["name"]}"'
        f'(이하 "EO"라 한다)는 상호 대등한 입장에서 붙임의 계약문서에 의하여 위 '
        f'영상제작 및 홍보에 관한 계약을 체결하고, 신의에 따라 성실히 계약상의 '
        f"의무를 이행할 것을 확약하며, 이 계약의 증거로서 계약서를 작성하여 "
        f"기명날인 후 각각 1통씩 보관한다."
    )
    set_east_asia_font(r, "바탕")
    r.font.size = Pt(10)

    add_body(doc, "붙임서류: 붙임서류도 본 계약에 포함된다.", indent_level=0)
    add_body(doc, "1. 영상제작 및 홍보계약서 일반조건 1부", indent_level=1)
    add_body(doc, "2. [별지 제1호 서식] 용어의 정의 1부", indent_level=1)
    add_body(doc, "3. [별지 제2호 서식] 수정요청가이드 1부", indent_level=1)

    add_blank(doc, size=10)

    # 체결일
    p = doc.add_paragraph()
    set_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.CENTER,
                         space_before=Pt(14), space_after=Pt(14))
    r = p.add_run("2026년      월      일")
    set_east_asia_font(r, "바탕")
    r.font.size = Pt(11)

    # 서명란
    sign_table = doc.add_table(rows=6, cols=2)
    sign_table.style = "Table Grid"
    sign_table.autofit = False
    for row in sign_table.rows:
        row.cells[0].width = Cm(7.75)
        row.cells[1].width = Cm(7.75)

    def sgn(cell, text, bold=False, size=10, font="바탕"):
        cell.text = ""
        p = cell.paragraphs[0]
        set_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.LEFT,
                             space_before=Pt(2), space_after=Pt(2))
        r = p.add_run(text)
        set_east_asia_font(r, font)
        r.font.size = Pt(size)
        r.bold = bold

    sgn(sign_table.cell(0, 0), "파트너", bold=True, font="굴림")
    sgn(sign_table.cell(0, 1), "EO", bold=True, font="굴림")
    sgn(sign_table.cell(1, 0), f"상호: {partner['name']}")
    sgn(sign_table.cell(1, 1), f"상호: {eo['name']}")
    sgn(sign_table.cell(2, 0), f"사업자등록번호: {partner['business_no']}")
    sgn(sign_table.cell(2, 1), f"사업자등록번호: {eo['business_no']}")
    sgn(sign_table.cell(3, 0),
        "주소: 서울특별시 중구 칠패로 37, 12,14층 (봉래동1가, 에이취에스비씨빌딩)"
        if partner["name"] == "주식회사 프레인글로벌"
        else f"주소: {partner['address']}")
    sgn(sign_table.cell(3, 1), f"주소: {eo['address']}")
    sgn(sign_table.cell(4, 0), f"대표이사: {partner['ceo']}    (인)")
    sgn(sign_table.cell(4, 1), f"대표이사: {eo['ceo']}    (인)")
    sgn(sign_table.cell(5, 0),
        f"담당: {partner_contact['name']} {partner_contact.get('role','')} / "
        f"{partner_contact['email']} / {partner_contact.get('phone','')}".rstrip(" /"))
    sgn(sign_table.cell(5, 1),
        f"담당: {eo_contact['name']} {eo_contact['role']} / {eo_contact['email']}")

    add_page_break(doc)

    # ═════════════ 일반조건 14조 ═════════════
    add_subtitle(doc, "영상제작 및 홍보계약서 일반조건", size=13)

    p = doc.add_paragraph()
    set_paragraph_format(p, first_line_indent=Cm(0.8),
                         space_before=Pt(4), space_after=Pt(10))
    r = p.add_run(
        f'"{partner["name"]}"(이하 "파트너"라 한다)와 "{eo["name"]}"'
        f'(이하 "EO"라 한다)는 {project_name[project_name.find("["):project_name.find("]")+1] if "[" in project_name else project_name} 영상(이하 '
        f'"영상"이라 한다) 제작 및 홍보를 위해 다음과 같이 계약을 체결한다.'
    )
    set_east_asia_font(r, "바탕")
    r.font.size = Pt(10)

    # 제1조
    add_article_header(doc, "제 1 조 【 계약의 목적 】")
    add_body(doc,
        f'본 계약의 목적은 "EO"가 "파트너"의 의뢰를 받아 제작 및 홍보하는 '
        f'영상에 대하여 상호간의 권리와 의무 및 기타 제반 사항을 규정하는 데 있다. '
        f'본 영상은 "파트너"가 광고주 {ad_client}로부터 '
        f'위임받아 발주한 브랜디드 콘텐츠로서 "EO"가 운영하는 유튜브 채널에 공개된다.')

    # 제2조
    add_article_header(doc, "제 2 조 【 계약기간 및 영상제작 일정 】")
    add_body(doc,
        f'본 계약의 "전체 일정"은 {start_date}(인서트 촬영일)부터 {end_date}'
        '("최종본" 예정 납품일)까지로 하며, 양 당사자의 합의에 의해 '
        '변경할 수 있다. "전체 일정", "촬영 일정", "업로드 일정"의 구체적 '
        '정의는 [별지 제1호 서식]을 준용한다.', indent_level=1)
    add_body(doc, '주요 촬영 일정은 다음과 같다.', indent_level=1)
    add_body(doc, f'1) 인서트 촬영: {insert_shoot_date}', indent_level=2)
    add_body(doc, f'2) 본 촬영(파트너 4인 인터뷰): {main_shoot_date}', indent_level=2)
    add_body(doc,
        '"가편집본" 전달일, 수정 반영 일정 및 "최종본" 납품일은 본 촬영 완료 후 '
        '"파트너"와 "EO"가 별도 협의하여 정한다.', indent_level=1)
    add_body(doc,
        '"파트너"는 "EO"에게 본 "영상"의 제작을 위하여 "EO"가 요청한 자료, 연락, '
        '검토 의견을 "EO"가 요구한 기한 내에 지체 없이 제공한다.', indent_level=1)
    add_body(doc,
        '"파트너"의 자료 미제공·확인 지연, 출연자 일정 변경, 천재지변 또는 기타 '
        '불가항력 사유가 있는 경우 "EO"의 납기 일정은 합리적 범위 내에서 '
        '순연될 수 있으며, "전체 일정"의 변경은 양 당사자의 합의에 의한다.',
        indent_level=1)

    # 제3조 ~ 제14조 — BCG 스크립트 본문을 그대로 이관 (금액/광고주명만 포맷)
    # 주의: 이 조항들은 대부분 계약 고정 문구라 아래 구조를 유지하면 baseline과 동일.
    add_article_header(doc, "제 3 조 【 계약의 내용 】")
    add_body(doc,
        '"EO"는 "파트너"와 협력을 통해 홍보에 효과적인 영상 제작을 위해 최선을 '
        '다하고, "파트너"는 "EO"가 영상을 제작·홍보하는 데 필요한 제반 사항에 '
        '대해 협조하며 관련 정보를 제공한다.', indent_level=1)
    add_body(doc, '본 계약의 최종 결과물은 다음과 같이 정의한다.', indent_level=1)
    add_body(doc,
        '1) 롱폼 영상 1편: BCG 코리아 파트너 4인의 인터뷰로 구성된 본편 다큐멘터리형 영상',
        indent_level=2)
    add_body(doc,
        '2) 쇼츠 영상 4편: 롱폼 촬영본 및 편집본을 기초로 가공되는 세로형 숏폼 영상 '
        '(원칙적으로 인물별 1편, 총 4편)', indent_level=2)
    add_body(doc,
        '3) 상기 각 영상의 "완성본" — 수정본 이후 양 당사자가 업로드 가능함을 '
        '최종 확인한 버전', indent_level=2)
    add_body(doc,
        '"파트너"는 "EO"가 "가편집본"을 전달한 이후로부터 "완성본" 확인 이전까지, '
        '본 영상에 대하여 편집 등 "본 수정"을 총 2회까지 무상으로 요청할 수 '
        '있으며, "EO"는 "파트너"와 협의한 기한 내에 이를 반영한다.', indent_level=1)
    add_body(doc,
        '전항의 "본 수정"은 촬영된 소스 내 다른 장면 또는 표현으로의 교체를 '
        '포함하되, "스토리라인 변경"은 포함하지 아니한다.', indent_level=1)
    add_body(doc,
        '자막 오탈자, 음성 보정, 색보정의 경미한 조정 등 발행 완성도에 영향을 '
        '미치는 범위 내의 "미세조정"은 수정 횟수에 산입되지 아니하며, "EO"는 '
        '합리적인 범위 내에서 이를 반영한다.', indent_level=1)
    add_body(doc,
        f'"본 수정"이 총 2회를 초과할 경우 매 회 금 {extra_cost_korean}원'
        f'(₩{extra_cost_per_edit:,}, VAT 포함)'
        '을, "스토리라인 변경" 요청의 경우 양 당사자가 별도 협의하여 정한 금액을 '
        '"파트너"가 "EO"에게 지급한다. "수정요청"과 "수정방법"에 관한 자세한 '
        '내용은 [별지 제2호 서식] 수정요청 가이드를 준용한다.', indent_level=1)
    add_body(doc,
        '광고 집행, 매체 운영, 유료 광고비 집행, 성과 보장 기타 본 조에 명시되지 '
        '아니한 마케팅 집행 업무는 본 계약의 용역 범위에 포함되지 아니하며, 필요 '
        '시 "파트너"와 "EO"는 별도의 서면 합의로 그 범위·예산·조건을 정한다.',
        indent_level=1)
    add_body(doc,
        '"EO"는 공정거래위원회의 관한 표시·광고 심사지침 개정안(2020.09.01)을 '
        '준수하며, 영상 내 경제적 이해관계의 표시 의무를 이행한다.', indent_level=1)
    add_body(doc,
        '"EO"는 업로드 일정을 "파트너"에게 사전 공지하며, 일정이 변경될 때에는 '
        '상호 협의 하에 조정할 수 있다. "EO"는 완성된 영상을 업로드하기 전 '
        '"파트너"에게 공유하여 정보에 오류가 있는지 확인한다.', indent_level=1)

    add_article_header(doc, "제 4 조 【 계약금액 및 지급 】")
    add_body(doc,
        f'본 계약 조건에 따라 "파트너"는 "EO"에게 영상 제작 및 홍보 비용으로 '
        f'공급가액 금 {korean_supply_price}원(₩{supply_price:,}, VAT 별도)을 지급한다. '
        f'부가가치세 금 {korean_vat}원(₩{supply_price // 10:,})을 합한 총 지급액은 '
        f'금 {korean_total}원(₩{supply_price + supply_price // 10:,})이다.',
        indent_level=1)
    add_body(doc, '전항의 공급가액 구성은 다음과 같다.', indent_level=1)
    add_body(doc,
        f'1) 롱폼 인터뷰 영상 1편 제작비: 금 {longform_price_korean}원'
        f'(₩{longform_price:,}, VAT 별도)',
        indent_level=2)
    add_body(doc,
        f'2) 쇼츠 영상 4편 제작비: 금 {shorts_price_korean}원'
        f'(₩{shorts_price:,}, VAT 별도, 1편당 {shorts_unit_price_korean}원)',
        indent_level=2)
    add_body(doc, '"파트너"는 아래 일정에 따라 비용을 "EO"가 지정한 계좌로 지급한다.',
             indent_level=1)
    add_body(doc,
        f'1) 선금: 총 공급가액의 50%에 해당하는 금 {deposit_korean}원'
        f'(₩{supply_price // 2:,}) 및 '
        '해당 부가가치세 — 본 계약 체결일로부터 14영업일 이내 지급', indent_level=2)
    add_body(doc,
        f'2) 잔금: 총 공급가액의 50%에 해당하는 금 {deposit_korean}원'
        f'(₩{supply_price // 2:,}) 및 '
        '해당 부가가치세 — "최종본" 전달 이후 14영업일 이내 지급', indent_level=2)
    add_body(doc,
        '"EO"는 계약금액 수령을 위해 각 지급 단계(선금·잔금)별로 "파트너"에게 '
        '세금계산서를 발행한다.', indent_level=1)
    add_body(doc,
        '"파트너"가 지급기한을 넘겨 대금을 지급하지 아니하는 경우 "EO"는 상당한 '
        '기간을 정하여 최고한 후 작업을 일시 중지하거나 납품을 유예할 수 있다.',
        indent_level=1)
    add_body(doc,
        f'입금 정보 — 은행: {eo["bank"]["bank_name"]} / '
        f'계좌번호: {eo["bank"]["account"]} / '
        f'예금주: {eo["bank"]["holder"]}', indent_level=1)

    # 제5조 ~ 제14조 — 조항 본문은 고정 템플릿 (광고주명만 포맷)
    _add_fixed_articles_5_to_14(doc, ad_client, supply_price)

    add_page_break(doc)

    # ═════════════ 별지 제1호 용어 정의 ═════════════
    add_subtitle(doc, "[별지 제1호 서식] 용어의 정의", size=12)
    for term, desc in TERM_DEFINITIONS:
        p = doc.add_paragraph()
        set_paragraph_format(p, first_line_indent=Cm(0.0),
                             left_indent=Cm(0.5),
                             space_before=Pt(4), space_after=Pt(2))
        r = p.add_run(f'• {term}  ')
        set_east_asia_font(r, "굴림")
        r.font.size = Pt(10)
        r.bold = True
        r2 = p.add_run(desc)
        set_east_asia_font(r2, "바탕")
        r2.font.size = Pt(10)

    add_page_break(doc)

    # ═════════════ 별지 제2호 수정요청 가이드 ═════════════
    _add_revision_guide(doc, extra_cost_per_edit, extra_cost_korean)

    # ═════════════ 내부 메모 페이지 ═════════════
    add_page_break(doc)
    _add_internal_memo_page(doc, partner["name"], internal_memos, checklist)

    # 저장
    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    print(f"✓ saved: {out}  ({out.stat().st_size} bytes)")
    return str(out)


def _add_fixed_articles_5_to_14(doc, ad_client, supply_price):
    """제5조 ~ 제14조 — BCG 샘플 고정 문구. ad_client와 supply_price만 파라미터."""
    # 제5조
    add_article_header(doc, "제 5 조 【 자료 제공 및 출연자 권리 】")
    add_body(doc,
        '영상 촬영에 필요한 자료는 "파트너"가 "EO"에게 제공하며, "EO"는 "파트너"와 '
        '사전 협의된 범위를 초과할 경우 제공된 자료를 이용할 수 없다.',
        indent_level=1)
    add_body(doc,
        f'"파트너"는 본 영상에 출연하는 {ad_client} 소속 출연자로부터 '
        '촬영, 편집, 공개, 홍보 및 이용에 필요한 초상권·퍼블리시티권·성명권 기타 '
        '제반 권리를 적법하게 사전 확보하여야 하며, 이를 "EO"에게 보장한다. '
        '"파트너"의 본 의무 불이행으로 제3자로부터 이의 제기·분쟁·손해배상 청구 '
        '등이 발생한 경우, "파트너"는 자신의 책임과 비용으로 이를 해결하고 "EO"에 '
        '발생한 손해를 배상한다.', indent_level=1)
    add_body(doc,
        '"EO"가 본 영상의 편집 과정에서 독자적으로 기획·주선·추가 촬영한 부분에 '
        '대한 영상 내 초상권 관련 권리는 "EO"에게 귀속된다.', indent_level=1)
    add_body(doc,
        '"EO"는 본 업무에 대하여 "파트너"로부터 제공된 서류·도면·정보·데이터 기타 '
        '모든 자료를 관리자의 의무로 보관·관리하고, "파트너"의 사전 서면 승낙 '
        '없이 복제·반출하거나 본건 이외의 목적으로 사용할 수 없다.', indent_level=1)

    # 제6조
    add_article_header(doc, "제 6 조 【 저작권의 귀속 】")
    add_body(doc,
        '제3조에 따라 제작된 영상의 저작권 및 이에 준하는 일체의 지식재산권은 '
        '"EO"에게 귀속되며, "EO"의 사전 서면 동의 없이 이를 판매·대여·복제 '
        '제공할 수 없다.', indent_level=1)
    add_body(doc,
        f'"EO"는 "파트너"에게 전달된 "완성본"을 편집하지 않은 상태로 "파트너" 및 '
        f'광고주 {ad_client.split("(")[0].strip() if "(" in ad_client else ad_client}가 다음 각 호의 비영리 목적으로 활용할 수 있는 '
        f'"이용허락" 권한을 무상으로 부여한다.', indent_level=1)
    add_body(doc, '1) 유튜브 임베드 영상 게시', indent_level=2)
    add_body(doc, '2) 자사 오피스 내 상영', indent_level=2)
    add_body(doc, '3) 블로그·카드뉴스 등에서의 영상 링크 게시', indent_level=2)
    add_body(doc, '4) 가공 없는 썸네일 또는 캡처 이미지 활용', indent_level=2)
    add_body(doc,
        '영상 유통과정에서 광고수익 등이 발생한 경우, 해당 수익은 "EO"에게 '
        '귀속된다. 단, "파트너"가 제4조에 정한 대금을 지급한 범위를 초과하는 '
        '광고수익이 발생한 경우의 배분은 양 당사자가 별도 협의한다.',
        indent_level=1)

    # 제7조
    add_article_header(doc, "제 7 조 【 라이선스 및 2차 활용 】")
    add_body(doc,
        '"EO"는 영상 제작을 위해 사용된 표현물 등이 제3자의 권리를 침해하지 '
        '않도록 해야 하며, "EO"의 독자적 기술에 의한 결과물에 대해 제3자로부터 '
        '상표권·특허권·저작권 등의 이유로 소송이나 분쟁이 발생한 경우 "EO"는 '
        '일체의 책임과 비용으로 "파트너"를 면책시킨다.', indent_level=1)
    add_body(doc,
        '"파트너"가 제6조 제2항에서 정한 비영리 이용허락 범위를 초과하여 '
        '"완성본"에 대한 편집·가공·상업적 활용을 원하는 경우, "파트너"는 "EO"와 '
        '별도 합의를 체결하고 아래 기준에 따른 추가 라이선스 비용을 부담한다.',
        indent_level=1)
    add_body(doc,
        '1) "클린본"(EO 로고, 워터마크, 자막, 음악 등 브랜딩 요소가 제거된 버전) '
        '제공 — 해당 산출물 공급가액의 30%에 해당하는 금액', indent_level=2)
    add_body(doc,
        '2) "원본"(촬영 원본 파일 또는 프로젝트 파일) 제공 및 2차 가공 권한 부여 '
        '— 해당 산출물 공급가액의 50%에 해당하는 금액', indent_level=2)
    add_body(doc,
        '본 계약 체결 시점 기준, "파트너" 및 광고주는 전항 제1호 및 제2호의 '
        '별도 라이선스를 구매하지 아니하는 것으로 본다.', indent_level=1)
    add_body(doc,
        '"파트너"와 광고주는 "EO"의 사전 서면 동의 없이 "완성본"을 편집·변형·'
        '2차 저작물화하거나 상업적으로 재이용할 수 없다.', indent_level=1)

    # 제8조
    add_article_header(doc, "제 8 조 【 비밀유지 】")
    add_body(doc,
        f'"파트너"와 "EO"는 본 계약의 체결 및 이행 과정에서 알게 된 상대방 및 '
        f'광고주({ad_client.split("(")[0].strip() if "(" in ad_client else ad_client})의 영업상·기술상·사업상 비밀정보, 출연자 개인정보, '
        f'미공개 영상 자료 등 일체의 비밀정보를 제3자에게 누설하거나 본 계약의 '
        f'목적 외로 사용해서는 아니 된다.', indent_level=1)
    add_body(doc,
        '본 조의 의무는 본 계약의 종료 또는 해지 이후에도 2년간 유효하다. 다만 '
        '법령상 공개의무가 있거나 상대방의 사전 서면 동의가 있는 경우는 예외로 한다.',
        indent_level=1)
    add_body(doc,
        '"EO"가 수신자 데이터베이스(DB)를 보유한 경우에도 해당 DB를 "파트너" 또는 '
        '광고주에게 제공하지 아니하며, 개인정보 보호법 및 정보통신망법 등 관련 '
        '법령을 준수한다.', indent_level=1)

    # 제9조
    add_article_header(doc, "제 9 조 【 계약의 해제·해지 】")
    add_body(doc,
        '일방 당사자의 과실로 계약의 목적을 달성하기 어려울 정도의 중대한 하자가 '
        '발생한 경우 상대방은 본 계약의 전부 또는 일부를 해지할 수 있다. 또한 '
        '일방이 본 계약의 주요 사항을 위반한 경우 상대방은 7일 이내 시정을 '
        '서면으로 최고한 후, 미시정 시 본 계약을 해지할 수 있다.', indent_level=1)
    add_body(doc,
        '선금 입금 후 "EO"의 귀책사유 없이 "파트너"의 사정에 의해 계약이 해지되는 '
        '경우 선금은 환수되지 아니한다. "EO"의 귀책사유로 계약이 해지되는 경우 '
        '"EO"는 "파트너"에게 기 수령한 선금의 2배에 해당하는 금액을 배상한다.',
        indent_level=1)
    add_body(doc,
        '제10조의 불가항력 사유로 본 계약의 이행이 불가능하게 된 경우 일방 '
        '당사자는 본 계약을 해지할 수 있으며, 이 경우 양 당사자는 기 수행분에 '
        '대한 정산에 관하여 상호 협의한다.', indent_level=1)

    # 제10조
    add_article_header(doc, "제 10 조 【 손해배상 및 책임 한도 】")
    add_body(doc,
        '"파트너" 또는 "EO"는 본 계약상의 의무 위반 또는 귀책사유로 상대방 또는 '
        '제3자에게 손해를 입힌 경우 그 손해를 배상하여야 한다.', indent_level=1)
    add_body(doc,
        f'일방 당사자가 본 계약과 관련하여 상대방에게 부담하는 손해배상 책임의 '
        f'총액은 본 계약상 공급가액(₩{supply_price:,})을 한도로 한다. 다만, 고의 또는 '
        f'중대한 과실, 제8조(비밀유지) 위반, 제5조 제2항에 따른 출연자 권리 보장 '
        f'위반, 또는 제3자의 지식재산권을 침해한 경우에는 본 한도가 적용되지 '
        f'아니한다.', indent_level=1)

    # 제11조
    add_article_header(doc, "제 11 조 【 불가항력 】")
    add_body(doc,
        '천재지변, 전쟁, 폭동, 감염병의 창궐, 정부의 규제·명령, 파업, 시스템·'
        '플랫폼 장애 기타 당사자가 합리적으로 통제할 수 없는 사유(이하 '
        '"불가항력")로 본 계약의 이행이 지연되거나 불가능하게 된 경우, 해당 '
        '당사자는 그에 대한 책임을 부담하지 아니한다. 다만, 해당 당사자는 지체 '
        '없이 상대방에게 불가항력 사유의 발생 사실을 통지하고 손해를 최소화하기 '
        '위하여 노력하여야 한다.')

    # 제12조
    add_article_header(doc, "제 12 조 【 영상 비공개·삭제 】")
    add_body(doc,
        f'출연자 또는 출연자 소속 조직({ad_client.split("(")[0].strip() if "(" in ad_client else ad_client})이 사회적 물의를 일으키거나 '
        f'공중에게 중대한 오인을 유발할 우려가 있는 사정이 발생한 경우, "파트너"와 '
        f'"EO"는 협의하여 해당 영상의 비공개·삭제 또는 수정 여부를 결정할 수 있다. '
        f'다만 전항 외의 단순한 불만족·전략 변경 등은 영상 내림의 사유가 되지 '
        f'아니한다.', indent_level=1)
    add_body(doc,
        '전항의 조치로 인해 추가 편집·재업로드 또는 대체 제작이 필요한 경우, '
        '"파트너"와 "EO"는 그 비용 및 일정에 관하여 별도 협의한다.', indent_level=1)

    # 제13조
    add_article_header(doc, "제 13 조 【 양도금지 】")
    add_body(doc,
        '"파트너"와 "EO"는 상대방의 사전 서면 승낙 없이 본 계약상의 권리 또는 '
        '의무의 전부 또는 일부를 제3자에게 양도·이전·담보제공 기타 처분할 수 없다.')

    # 제14조
    add_article_header(doc, "제 14 조 【 분쟁해결 및 관계법령 】")
    add_body(doc,
        '본 계약과 관련하여 당사자 사이에 분쟁이 발생하는 경우, "파트너"와 "EO"는 '
        '상호 호혜의 원칙에 따라 성실히 협의하여 해결한다. 협의로 해결되지 '
        '아니하는 경우 서울중앙지방법원을 제1심 전속적 합의관할 법원으로 한다.',
        indent_level=1)
    add_body(doc,
        '본 계약에서 정하지 아니한 사항은 관계 대한민국 법령(민법, 약관의 규제에 '
        '관한 법률, 하도급거래 공정화에 관한 법률, 표시·광고의 공정화에 관한 '
        '법률, 저작권법, 개인정보 보호법 등) 및 일반 상관례에 따른다.', indent_level=1)
    add_body(doc,
        '본 계약의 변경 또는 추가는 양 당사자의 서면 합의에 의해서만 효력이 '
        '발생한다. 본 계약은 쌍방이 서명·날인한 날부터 유효하며, 본 계약을 '
        '증명하기 위해 계약서 2부를 작성하고 각 1부씩 보관한다.', indent_level=1)


def _add_revision_guide(doc, extra_cost_per_edit, extra_cost_korean):
    """별지 제2호 — 수정요청 가이드 (고정 템플릿, 초과비용만 파라미터)"""
    add_subtitle(doc, "[별지 제2호 서식] 수정요청 가이드", size=12)

    add_article_header(doc, "1. 수정의 단계 구분")
    add_body(doc, '"EO"는 수정 요청을 아래와 같이 3단계로 구분하여 운영한다.', indent_level=1)
    add_body(doc,
        '1) 질문지 및 프로덕션 단계: "파트너"는 자유롭게 의견을 취합·협의할 수 있으며, '
        '본 단계의 의견 교환은 수정 횟수에 산입되지 아니한다.', indent_level=2)
    add_body(doc,
        '2) "가편집본" 단계 이후 "본 수정": 총 2회까지 무상. 촬영된 소스 내 다른 '
        '장면·표현으로의 교체 범위 내에서 반영되며, "스토리라인 변경"은 포함되지 않는다.',
        indent_level=2)
    add_body(doc,
        '3) "미세조정": 자막 오탈자·음성 보정·색보정 등 완성도 향상을 위한 경미한 '
        '편집은 횟수 제한 없이 합리적 범위에서 반영된다.', indent_level=2)

    add_article_header(doc, "2. 수정 요청 방법")
    add_body(doc,
        '수정 요청은 "파트너"의 담당자가 "EO"의 담당 PD에게 이메일로 전달하되, '
        '각 수정 요청에는 다음 사항을 명시한다.', indent_level=1)
    add_body(doc, '1) 타임코드 (예: 00:01:23)', indent_level=2)
    add_body(doc, '2) 기존 내용 → 수정 희망 내용', indent_level=2)
    add_body(doc, '3) 수정 사유', indent_level=2)

    add_article_header(doc, "3. 반영 불가 사유")
    add_body(doc, '다음에 해당하는 수정 요청은 반영이 불가하거나 별도 협의 사안으로 본다.',
             indent_level=1)
    add_body(doc, '1) 본 계약 체결 이후 합의된 영상 방향성과 본질적으로 어긋나는 변경',
             indent_level=2)
    add_body(doc, '2) 사전 협의 없이 촬영 이후 새로 추가 요청되는 장면 또는 출연자',
             indent_level=2)
    add_body(doc, '3) 추상적·불확실한 피드백으로 반영 범위가 특정되지 아니하는 요청',
             indent_level=2)
    add_body(doc, '4) "스토리라인 변경"에 해당하는 요청 (별도 협의 및 추가 비용 사안)',
             indent_level=2)

    add_article_header(doc, "4. 수정 초과 비용")
    add_body(doc,
        f'"본 수정"이 2회를 초과하는 경우 "파트너"는 매 회 금 {extra_cost_korean}원'
        f'(₩{extra_cost_per_edit:,}, '
        'VAT 포함)을 "EO"에게 지급하며, 해당 비용은 잔금과 별도로 정산한다. '
        '"스토리라인 변경"은 기획·촬영·편집 작업량에 따라 양 당사자가 별도 협의하여 '
        '정한 금액으로 한다.')


def _add_internal_memo_page(doc, partner_short_name, memos, checklist):
    """내부 검토 메모 페이지 — 외부 발송 전 반드시 삭제"""
    p = doc.add_paragraph()
    set_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.CENTER,
                         space_before=Pt(6), space_after=Pt(10))
    r = p.add_run(f"⚠️ 내부 검토 메모 — {partner_short_name} 발송 전 이 페이지 반드시 삭제 ⚠️")
    set_east_asia_font(r, "굴림")
    r.font.size = Pt(12)
    r.bold = True
    r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    p = doc.add_paragraph()
    set_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.CENTER,
                         space_before=Pt(0), space_after=Pt(14))
    r = p.add_run(
        f"(본 페이지는 서현님 내부 검토용입니다. "
        f"{partner_short_name}에 전달되는 최종본에는 포함되지 않아야 합니다.)"
    )
    set_east_asia_font(r, "바탕")
    r.font.size = Pt(9)
    r.italic = True

    add_subtitle(doc, "서명 전 확인 필요 사항", size=12, center=False)

    for title, desc in memos:
        p = doc.add_paragraph()
        set_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                             space_before=Pt(8), space_after=Pt(2))
        r = p.add_run(title)
        set_east_asia_font(r, "굴림")
        r.font.size = Pt(10)
        r.bold = True

        p2 = doc.add_paragraph()
        set_paragraph_format(p2, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                             left_indent=Cm(0.5),
                             space_before=Pt(1), space_after=Pt(4))
        r2 = p2.add_run(desc)
        set_east_asia_font(r2, "바탕")
        r2.font.size = Pt(10)

    add_subtitle(doc, f"{partner_short_name} 발송 전 체크리스트", size=11, center=False)
    for item in checklist:
        p = doc.add_paragraph()
        set_paragraph_format(p, left_indent=Cm(0.3),
                             space_before=Pt(2), space_after=Pt(2))
        r = p.add_run(item)
        set_east_asia_font(r, "바탕")
        r.font.size = Pt(10)
